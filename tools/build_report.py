from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSIGNMENT_ROOT = ROOT.parent
SCREENSHOTS = ROOT / "artifacts" / "screenshots"
REPORT_NAME = "Assignment_4_File_System_Simulation_Report.docx"
HOSTED_URL = "https://os-file-system-simulator-assignment.vercel.app"
REPO_URL = "https://github.com/babarnaeem0001/os-file-system-simulator-assignment-4"

PURPLE = "9333EA"
PURPLE_DARK = "581C87"
PURPLE_LIGHT = "F5EDFF"
INK = "111827"
MUTED = "4B5563"
LINE = "E7E2EF"
GREEN_LIGHT = "ECFDF5"
AMBER_LIGHT = "FFFBEB"
RED_LIGHT = "FEF2F2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color=LINE):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    tc_pr.append(borders)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = margins.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_run(run, size=None, bold=None, color=None, font="Aptos"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def paragraph_text(paragraph, text, size=10.5, color=INK, bold=False, font="Aptos"):
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold, color=color, font=font)
    return run


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), PURPLE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_section_heading(doc, title, kicker=None):
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(0)
        paragraph_text(p, kicker.upper(), size=8.5, color=PURPLE, bold=True)
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(8)
    paragraph_text(heading, title, size=18, color=INK, bold=True)


def add_callout(doc, title, body, fill=PURPLE_LIGHT, accent=PURPLE_DARK):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "D8B4FE")
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p = cell.paragraphs[0]
    paragraph_text(p, title, size=11, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    paragraph_text(p2, body, size=10.5, color=INK)
    doc.add_paragraph()


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(4.7)
    for idx, (key, value) in enumerate(rows):
        key_cell = table.cell(idx, 0)
        value_cell = table.cell(idx, 1)
        set_cell_shading(key_cell, "F7F5FB")
        set_cell_border(key_cell)
        set_cell_border(value_cell)
        set_cell_margins(key_cell)
        set_cell_margins(value_cell)
        key_p = key_cell.paragraphs[0]
        paragraph_text(key_p, key, size=9.5, color=MUTED, bold=True)
        value_p = value_cell.paragraphs[0]
        if isinstance(value, tuple):
            label, url = value
            add_hyperlink(value_p, label, url)
        else:
            paragraph_text(value_p, str(value), size=10.5, color=INK, bold=True)
    doc.add_paragraph()


def add_feature_table(doc):
    rows = [
        ("File creation, deletion, reading, writing", "Implemented", "CRUD controls and route-backed actions."),
        ("Contiguous allocation", "Implemented", "Finds a free run and stores adjacent data blocks."),
        ("Linked allocation", "Implemented", "Places data blocks flexibly and models traversal cost."),
        ("Indexed allocation", "Implemented", "Allocates an index block plus data blocks."),
        ("File system interface for processes", "Implemented", "Each operation is issued by one of the four team processes."),
        ("Disk usage and fragmentation metrics", "Implemented", "Calculates used blocks, free runs, largest run, and fragmentation."),
        ("File access times", "Implemented", "Estimates time based on allocation method, block count, mode, and fragmentation."),
        ("Shared-file conflict and deadlock", "Implemented", "Read/write locks, wait-for graph cycles, and recovery button."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    widths = [Inches(2.8), Inches(1.35), Inches(2.75)]
    headers = ("Requirement", "Status", "Evidence")
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.width = widths[idx]
        set_cell_shading(cell, PURPLE)
        set_cell_border(cell, PURPLE)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        paragraph_text(p, header, size=9.5, color="FFFFFF", bold=True)
    for requirement, status, evidence in rows:
        row = table.add_row()
        values = (requirement, status, evidence)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            cell.width = widths[idx]
            set_cell_border(cell)
            set_cell_margins(cell)
            if idx == 1:
                set_cell_shading(cell, GREEN_LIGHT)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph_text(cell.paragraphs[0], value, size=9.5, color="047857", bold=True)
            else:
                paragraph_text(cell.paragraphs[0], value, size=9.5, color=INK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()


def add_team_table(doc):
    members = [
        ("Hamid Saleem", "9061", "Process P-9061"),
        ("Babar Naeem", "8963", "Process P-8963"),
        ("Muhammad Sabeel Khan", "8926", "Process P-8926"),
        ("Abdul Sami", "8929", "Process P-8929"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.autofit = False
    headers = ("Member", "Student ID", "Simulator Role")
    widths = [Inches(2.6), Inches(1.4), Inches(2.7)]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.width = widths[idx]
        set_cell_shading(cell, PURPLE_DARK)
        set_cell_border(cell, PURPLE_DARK)
        set_cell_margins(cell)
        paragraph_text(cell.paragraphs[0], header, size=9.5, color="FFFFFF", bold=True)
    for member, student_id, role in members:
        row = table.add_row()
        for idx, value in enumerate((member, student_id, role)):
            cell = row.cells[idx]
            cell.width = widths[idx]
            set_cell_border(cell)
            set_cell_margins(cell)
            paragraph_text(cell.paragraphs[0], value, size=10, color=INK, bold=idx == 0)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        paragraph_text(p, item, size=10.5, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        paragraph_text(p, item, size=10.5, color=INK)


def add_screenshot(doc, title, filename, caption, width=6.6):
    path = SCREENSHOTS / filename
    add_section_heading(doc, title, "Hosted screenshot")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_text(cap, caption, size=9, color=MUTED, bold=True)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ("List Bullet", "List Number"):
        styles[style_name].font.name = "Aptos"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        styles[style_name].font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_text(header, "CS313 Assignment 4 - OS File System Simulator", size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_text(footer, "GitHub + Vercel deployment evidence included", size=8.5, color=MUTED)

    return doc


def build_report():
    doc = setup_document()

    cover = doc.add_table(rows=1, cols=1)
    cell = cover.cell(0, 0)
    set_cell_shading(cell, PURPLE)
    set_cell_border(cell, PURPLE)
    set_cell_margins(cell, top=360, start=360, bottom=360, end=360)
    p = cell.paragraphs[0]
    paragraph_text(p, "Operating System Concepts", size=11, color="FFFFFF", bold=True)
    p2 = cell.add_paragraph()
    paragraph_text(p2, "OS File System Simulator", size=28, color="FFFFFF", bold=True)
    p3 = cell.add_paragraph()
    paragraph_text(p3, "Assignment 4: File System Simulation", size=14, color="F3E8FF", bold=True)

    doc.add_paragraph()
    add_key_value_table(
        doc,
        [
            ("Course Code", "CS313"),
            ("Course Title", "Operating System Concepts"),
            ("Instructor", "Ms. Naveen Ahmed"),
            ("University", "Abasyn University"),
            ("Semester", "Spring 2026"),
            ("Prepared For", "Assignment 4 only"),
            ("Hosted Website", ("os-file-system-simulator-assignment.vercel.app", HOSTED_URL)),
            ("GitHub Repository", ("babarnaeem0001/os-file-system-simulator-assignment-4", REPO_URL)),
        ],
    )

    add_section_heading(doc, "Team Members")
    add_team_table(doc)

    doc.add_page_break()
    add_section_heading(doc, "Executive Overview", "Project")
    add_callout(
        doc,
        "Submission Summary",
        "The project is a production-deployed Next.js application that simulates file system operations, allocation strategies, disk performance, process file access, synchronization, and deadlock recovery for Assignment 4.",
    )
    paragraph_text(
        doc.add_paragraph(),
        "The simulator focuses on the file system component of the operating system project. It gives users a dashboard where files can be created, read, written, deleted, allocated on a simulated disk, accessed by processes, and used to demonstrate shared-resource conflicts.",
    )
    add_bullets(
        doc,
        [
            "A 48-block simulated disk is used, with each block representing 4 KB.",
            "Files can be allocated through contiguous, linked, or indexed allocation.",
            "The frontend posts each operation to a Vercel route handler at /api/filesystem.",
            "The backend returns the next simulation state, including disk blocks, locks, wait queues, metrics, and deadlock cycles.",
            "The application is deployed through a GitHub-connected Vercel production project.",
        ],
    )

    add_section_heading(doc, "Requirement Coverage", "Assignment 4")
    add_feature_table(doc)

    add_section_heading(doc, "Architecture", "Implementation")
    paragraph_text(
        doc.add_paragraph(),
        "The application uses a stateless serverless backend pattern. The client sends the current state and a typed action to the route handler. The backend executes the simulation logic and returns a new state. This makes the app easy to host on Vercel because no external database is required.",
    )
    add_key_value_table(
        doc,
        [
            ("Frontend", "React client component dashboard with Tailwind CSS and lucide-react icons."),
            ("Backend", "Next.js route handler: GET/POST /api/filesystem."),
            ("Simulation Core", "TypeScript engine in src/lib/filesystem.ts."),
            ("Deployment", "GitHub repository connected to Vercel production hosting."),
            ("Domain", HOSTED_URL),
        ],
    )

    add_section_heading(doc, "File Allocation Algorithms", "Algorithms")
    add_numbered(
        doc,
        [
            "Contiguous allocation searches for a consecutive run of free blocks large enough to store the file. It is fast for sequential access but can fail when free space is fragmented.",
            "Linked allocation selects available blocks without requiring adjacency. It reduces allocation failure caused by fragmentation, but access time increases because blocks must be traversed in sequence.",
            "Indexed allocation reserves one index block and stores pointers to all data blocks. It gives direct lookup behavior while consuming one extra block for the index structure.",
        ],
    )
    add_callout(
        doc,
        "Performance Metrics",
        "The simulator calculates disk usage, free block count, number of free runs, largest free run, fragmentation percentage, average access time, active locks, and waiting requests after every operation.",
        fill=AMBER_LIGHT,
        accent="B45309",
    )

    add_section_heading(doc, "Synchronization and Deadlock", "Concurrency")
    paragraph_text(
        doc.add_paragraph(),
        "Every file access request is issued by a process. Read locks can be shared, while write locks require exclusive access. When a request cannot be granted, it enters the waiting queue and records the process that is blocking it.",
    )
    add_bullets(
        doc,
        [
            "Active locks store the file, mode, holder process, and tick when the lock was acquired.",
            "Waiting requests form a wait-for graph between blocked processes and lock holders.",
            "A cycle in the wait-for graph is reported as a deadlock.",
            "Deadlock recovery preempts one process, releases its locks, and retries waiting requests.",
        ],
    )
    add_callout(
        doc,
        "Demonstrated Scenario",
        "Hamid Saleem holds a write lock on ledger.db while Babar Naeem holds a write lock on shared.cfg. Each then requests the other's file, creating a two-process deadlock that the app detects and resolves.",
        fill=RED_LIGHT,
        accent="B91C1C",
    )

    add_section_heading(doc, "User Interface Design", "Design")
    paragraph_text(
        doc.add_paragraph(),
        "The interface follows the provided local design guidance: top-bar navigation, a clean dashboard layout, purple as the primary accent, glass-like panels, soft shadows, compact cards, and clear hierarchy. The UI avoids a sidebar and keeps the main workflow on the first screen.",
    )
    add_bullets(
        doc,
        [
            "Top navigation contains Overview, Disk, Files, and Conflicts views.",
            "Cards and tables are used for repeated data, metrics, process records, and file records.",
            "Buttons include icons for create, read, write, release, compact, reset, and resolve actions.",
            "Forms use visible labels, keyboard focus states, and accessible contrast.",
        ],
    )

    add_section_heading(doc, "Deployment and Verification", "Evidence")
    add_key_value_table(
        doc,
        [
            ("GitHub Repo", ("View repository", REPO_URL)),
            ("Production URL", ("Open hosted site", HOSTED_URL)),
            ("Vercel Target", "Production"),
            ("Deployment Status", "Ready"),
            ("API Verification", "GET /api/filesystem returned HTTP 200."),
            ("Browser Verification", "No console errors, no Next.js overlay, hosted pages loaded successfully."),
            ("Build Verification", "npm run lint and npm run build completed successfully."),
            ("Log Check", "Vercel error logs showed no errors for the last hour after deployment."),
        ],
    )

    doc.add_page_break()
    add_screenshot(
        doc,
        "Hosted Dashboard",
        "hosted-home.png",
        "Figure 1: Production overview page showing file operations, process access, metrics, and backend status.",
    )
    doc.add_page_break()
    add_screenshot(
        doc,
        "Disk Allocation View",
        "hosted-disk.png",
        "Figure 2: Disk block map showing contiguous, linked, and indexed allocation blocks.",
    )
    doc.add_page_break()
    add_screenshot(
        doc,
        "Deadlock Detection View",
        "hosted-conflicts.png",
        "Figure 3: Hosted conflict scenario with active locks, waiting queue, and detected deadlock cycle.",
    )
    doc.add_page_break()
    add_screenshot(
        doc,
        "Mobile Layout",
        "hosted-mobile-viewport.png",
        "Figure 4: Responsive mobile viewport of the deployed simulator.",
        width=2.35,
    )

    add_section_heading(doc, "Conclusion", "Result")
    paragraph_text(
        doc.add_paragraph(),
        "The Assignment 4 file system simulation is complete, deployed, and verified. It demonstrates file operations, allocation schemes, disk performance tracking, process file access, synchronization, and deadlock handling in one cohesive dashboard.",
    )

    parent_output = ASSIGNMENT_ROOT / REPORT_NAME
    repo_output = ROOT / "docs" / REPORT_NAME
    doc.save(parent_output)
    doc.save(repo_output)
    print(parent_output)
    print(repo_output)


if __name__ == "__main__":
    build_report()
